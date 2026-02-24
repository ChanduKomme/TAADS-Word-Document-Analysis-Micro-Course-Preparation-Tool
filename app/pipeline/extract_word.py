"""
Word document text extraction module.
Extracts text, paragraphs, and images from Word documents (.docx).
"""

from pathlib import Path
from typing import List, Dict, Any, Optional, Union
import re
from docx import Document
from docx.shared import Inches
import io


def extract_word_texts(
    docx_path: Union[str, Path],
    force_ocr: bool = False,
    use_ai: bool = False
) -> List[Dict[str, Any]]:
    """
    Extract text from a Word document, treating each page-worth of content as a "page".
    
    Since Word documents don't have explicit pages like PDFs, we estimate pages
    based on paragraph count or create logical sections.
    
    Args:
        docx_path: Path to the Word document
        force_ocr: Not applicable for Word (ignored)
        use_ai: Use AI enhancement (not implemented)
        
    Returns:
        List of dictionaries, one per logical "page", containing:
        - page: Page number (1-indexed)
        - text: Extracted text
        - word_count: Number of words
        - char_count: Number of characters
        - ocr_used: Always False for Word docs
        - quality_metrics: Quality assessment dictionary
    """
    docx_path = Path(docx_path)
    
    if not docx_path.exists():
        raise FileNotFoundError(f"Word document not found: {docx_path}")
    
    doc = Document(str(docx_path))
    
    all_paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            all_paragraphs.append(text)
    
    full_text = "\n\n".join(all_paragraphs)
    
    paragraphs_per_page = 15
    results = []
    
    if len(all_paragraphs) == 0:
        results.append({
            "page": 1,
            "text": "",
            "word_count": 0,
            "char_count": 0,
            "ocr_used": False,
            "quality_metrics": _calculate_quality_metrics("")
        })
    else:
        page_num = 1
        for i in range(0, len(all_paragraphs), paragraphs_per_page):
            page_paragraphs = all_paragraphs[i:i + paragraphs_per_page]
            page_text = "\n\n".join(page_paragraphs)
            
            word_count = len(page_text.split())
            char_count = len(page_text)
            
            results.append({
                "page": page_num,
                "text": page_text,
                "word_count": word_count,
                "char_count": char_count,
                "ocr_used": False,
                "quality_metrics": _calculate_quality_metrics(page_text)
            })
            page_num += 1
    
    return results


def _calculate_quality_metrics(text: str) -> Dict[str, Any]:
    """Calculate quality metrics for extracted text."""
    if not text:
        return {
            "word_count": 0,
            "char_count": 0,
            "avg_word_length": 0,
            "sentence_count": 0,
            "ocr_confidence": 1.0,
            "quality_score": 0
        }
    
    words = text.split()
    word_count = len(words)
    char_count = len(text)
    avg_word_length = sum(len(w) for w in words) / max(word_count, 1)
    sentence_count = len(re.findall(r'[.!?]+', text))
    
    valid_words = sum(1 for w in words if 2 <= len(w) <= 20)
    valid_word_ratio = valid_words / max(word_count, 1)
    
    quality_score = min(1.0, valid_word_ratio * 0.8 + 0.2)
    
    return {
        "word_count": word_count,
        "char_count": char_count,
        "avg_word_length": round(avg_word_length, 2),
        "sentence_count": sentence_count,
        "ocr_confidence": 1.0,
        "quality_score": round(quality_score, 3)
    }


def extract_word_images(
    docx_path: Union[str, Path],
    output_dir: Union[str, Path]
) -> List[Dict[str, Any]]:
    """
    Extract embedded images from a Word document.
    
    Args:
        docx_path: Path to the Word document
        output_dir: Directory to save extracted images
        
    Returns:
        List of dictionaries with image information
    """
    docx_path = Path(docx_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    doc = Document(str(docx_path))
    images = []
    
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
            image_data = rel.target_part.blob
            
            content_type = rel.target_part.content_type
            if "png" in content_type:
                ext = "png"
            elif "jpeg" in content_type or "jpg" in content_type:
                ext = "jpg"
            elif "gif" in content_type:
                ext = "gif"
            else:
                ext = "png"
            
            image_filename = f"image_{image_count:03d}.{ext}"
            image_path = output_dir / image_filename
            
            with open(image_path, "wb") as f:
                f.write(image_data)
            
            images.append({
                "page": 1,
                "seq": image_count,
                "image_path": str(image_path),
                "bbox": [0, 0, 100, 100],
                "area": 10000,
                "width": 100,
                "height": 100
            })
    
    return images


def extract_word_tables(docx_path: Union[str, Path]) -> List[Dict[str, Any]]:
    """
    Extract tables from a Word document.
    
    Args:
        docx_path: Path to the Word document
        
    Returns:
        List of dictionaries with table data
    """
    docx_path = Path(docx_path)
    doc = Document(str(docx_path))
    
    tables = []
    for idx, table in enumerate(doc.tables, start=1):
        rows_data = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(row_cells)
        
        if rows_data:
            tables.append({
                "page": 1,
                "seq": idx,
                "rows": len(rows_data),
                "cols": len(rows_data[0]) if rows_data else 0,
                "data": rows_data,
                "bbox": [0, 0, 100, 100]
            })
    
    return tables
